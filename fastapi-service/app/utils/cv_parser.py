"""
CV Parser – extracts text AND structured sections from PDF / DOCX.
Preserves section ordering. Uses pdfplumber (primary) + PyMuPDF (fallback).
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Dict, List, Optional, Tuple

# ── PDF ──────────────────────────────────────────────────────────────────────
try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    _HAS_PYMUPDF = True
except ImportError:
    _HAS_PYMUPDF = False

# ── DOCX ─────────────────────────────────────────────────────────────────────
try:
    import docx as python_docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

from ..schemas.cv import (
    StructuredCV, PersonalInformation, ExperienceEntry,
    EducationEntry, ProjectEntry, CertificationEntry,
)

# ── Section heading detection ────────────────────────────────────────────────
SECTION_PATTERNS: List[Tuple[str, str]] = [
    (r"(?i)\b(summary|profile|objective|about me|professional summary)\b", "summary"),
    (r"(?i)\b(experience|work experience|employment|professional experience|career)\b", "experience"),
    (r"(?i)\b(education|academic|qualifications?|studies)\b", "education"),
    (r"(?i)\b(skills?|technical skills?|core competencies|competencies|technologies)\b", "skills"),
    (r"(?i)\b(projects?|personal projects?|academic projects?|portfolio)\b", "projects"),
    (r"(?i)\b(certifications?|licenses?|courses?|training)\b", "certifications"),
    (r"(?i)\b(languages?)\b", "languages"),
    (r"(?i)\b(contact|personal info|personal information|details)\b", "contact"),
]

EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE    = re.compile(r"(\+?\d[\d\s\-\(\)]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
GITHUB_RE   = re.compile(r"github\.com/[\w\-]+", re.IGNORECASE)
URL_RE      = re.compile(r"https?://[^\s]+")
DATE_RE     = re.compile(
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"[\s,]*\d{4}|(?:\d{4})\s*[-–]\s*(?:\d{4}|Present|Current|Now)",
    re.IGNORECASE
)
BULLET_RE   = re.compile(r"^[\s•\-\*\u2022\u2023\u25E6\u2043]+")


def _detect_section(line: str) -> Optional[str]:
    stripped = line.strip()
    if len(stripped) > 60:
        return None
    for pattern, name in SECTION_PATTERNS:
        if re.search(pattern, stripped):
            return name
    return None


def _split_into_sections(text: str) -> Dict[str, str]:
    """Split raw text into ordered sections dict."""
    lines = text.splitlines()
    sections: Dict[str, List[str]] = {}
    current_section = "header"
    sections[current_section] = []

    for line in lines:
        detected = _detect_section(line)
        if detected and len(line.strip()) < 60:
            current_section = detected
            if current_section not in sections:
                sections[current_section] = []
        else:
            sections.setdefault(current_section, []).append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if any(l.strip() for l in v)}


def _parse_personal_info(header_text: str, full_text: str) -> PersonalInformation:
    search_text = header_text or full_text[:500]

    email_m    = EMAIL_RE.search(search_text)
    phone_m    = PHONE_RE.search(search_text)
    linkedin_m = LINKEDIN_RE.search(full_text)
    github_m   = GITHUB_RE.search(full_text)

    # name: first non-empty, non-header-keyword line
    name = None
    for line in search_text.splitlines():
        stripped = line.strip()
        if stripped and len(stripped) < 60 and not EMAIL_RE.search(stripped) \
                and not PHONE_RE.search(stripped) and not _detect_section(stripped):
            name = stripped
            break

    return PersonalInformation(
        full_name  = name,
        email      = email_m.group(0)    if email_m    else None,
        phone      = phone_m.group(0)    if phone_m    else None,
        linkedin   = linkedin_m.group(0) if linkedin_m else None,
        github     = github_m.group(0)   if github_m   else None,
    )


def _parse_experience(text: str) -> List[ExperienceEntry]:
    entries: List[ExperienceEntry] = []
    if not text:
        return entries

    # Split on blank lines between entries
    blocks = re.split(r"\n{2,}", text.strip())
    for block in blocks:
        lines = [l for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        entry = ExperienceEntry()
        bullets: List[str] = []
        for line in lines:
            stripped = line.strip()
            date_m = DATE_RE.search(stripped)
            
            # Extract date if found
            if date_m and not entry.start_date:
                date_str = date_m.group(0)
                if "–" in date_str or "-" in date_str:
                    parts = re.split(r"[-–]", date_str, 1)
                    entry.start_date = parts[0].strip()
                    entry.end_date   = parts[1].strip() if len(parts) > 1 else None
                else:
                    entry.start_date = date_str
                
                # Strip the date part from the line to check for job title/company info
                title_part = stripped.replace(date_str, "").strip(" (株),-–() ")
                if title_part and not entry.job_title and len(title_part) < 80 and not BULLET_RE.match(title_part):
                    if " at " in title_part:
                        t_parts = title_part.split(" at ", 1)
                        entry.job_title = t_parts[0].strip()
                        entry.company   = t_parts[1].strip()
                    elif " @ " in title_part:
                        t_parts = title_part.split(" @ ", 1)
                        entry.job_title = t_parts[0].strip()
                        entry.company   = t_parts[1].strip()
                    else:
                        entry.job_title = title_part
            elif BULLET_RE.match(stripped):
                bullets.append(BULLET_RE.sub("", stripped).strip())
            elif not entry.job_title and len(stripped) < 80:
                if " at " in stripped:
                    t_parts = stripped.split(" at ", 1)
                    entry.job_title = t_parts[0].strip()
                    entry.company   = t_parts[1].strip()
                elif " @ " in stripped:
                    t_parts = stripped.split(" @ ", 1)
                    entry.job_title = t_parts[0].strip()
                    entry.company   = t_parts[1].strip()
                else:
                    entry.job_title = stripped
            elif not entry.company and len(stripped) < 80:
                entry.company = stripped
        entry.bullets = bullets
        if entry.job_title or entry.company or bullets:
            entries.append(entry)
    return entries


def _parse_education(text: str) -> List[EducationEntry]:
    entries: List[EducationEntry] = []
    if not text:
        return entries
    blocks = re.split(r"\n{2,}", text.strip())
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        entry = EducationEntry()
        for line in lines:
            date_m = DATE_RE.search(line)
            if date_m and not entry.start_date:
                date_str = date_m.group(0)
                if "–" in date_str or "-" in date_str:
                    parts = re.split(r"[-–]", date_str, 1)
                    entry.start_date = parts[0].strip()
                    entry.end_date   = parts[1].strip() if len(parts) > 1 else None
                else:
                    entry.end_date = date_str
            elif re.search(r"(?i)\b(bachelor|master|phd|b\.?sc|m\.?sc|b\.?eng|associate|diploma)\b", line):
                entry.degree = line
            elif re.search(r"(?i)(university|college|institute|school|academy)", line):
                entry.institution = line
            elif re.search(r"(?i)gpa|cgpa", line):
                gpa_m = re.search(r"[\d.]+\s*/\s*[\d.]+|[\d.]+", line)
                entry.gpa = gpa_m.group(0) if gpa_m else line
        if entry.degree or entry.institution:
            entries.append(entry)
    return entries


def _parse_skills(text: str) -> List[str]:
    if not text:
        return []
    skills = []
    for line in text.splitlines():
        # split on commas, pipes, bullets
        parts = re.split(r"[,|•\-\n]+", line)
        for part in parts:
            s = part.strip()
            if s and len(s) < 50:
                skills.append(s)
    return list(dict.fromkeys(s for s in skills if s))


def _parse_projects(text: str) -> List[ProjectEntry]:
    entries: List[ProjectEntry] = []
    if not text:
        return entries
    blocks = re.split(r"\n{2,}", text.strip())
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        entry = ProjectEntry(name=lines[0] if lines else None)
        desc_lines = []
        for line in lines[1:]:
            url_m = URL_RE.search(line)
            if url_m:
                entry.url = url_m.group(0)
            elif BULLET_RE.match(line):
                desc_lines.append(BULLET_RE.sub("", line).strip())
            else:
                desc_lines.append(line)
        entry.description = " ".join(desc_lines) if desc_lines else None
        # extract tech tags (words in parens or after "Tech:" / "Technologies:")
        tech_m = re.search(r"(?i)(tech(?:nologies)?|stack|tools?)[:–\-]?\s*(.+)", block)
        if tech_m:
            entry.technologies = [t.strip() for t in re.split(r"[,|]", tech_m.group(2)) if t.strip()]
        if entry.name:
            entries.append(entry)
    return entries


def _parse_certifications(text: str) -> List[CertificationEntry]:
    entries: List[CertificationEntry] = []
    if not text:
        return entries
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            date_m = DATE_RE.search(stripped)
            cert_name = DATE_RE.sub("", stripped).strip(" •-")
            entries.append(CertificationEntry(
                name = cert_name or stripped,
                date = date_m.group(0) if date_m else None,
            ))
    return entries


def _parse_languages(text: str) -> List[str]:
    if not text:
        return []
    langs = []
    for part in re.split(r"[,|\n•\-]+", text):
        s = part.strip()
        if s and len(s) < 40:
            langs.append(s)
    return list(dict.fromkeys(langs))


class CVParser:
    # ── Text Extraction ───────────────────────────────────────────────────────

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        text = ""
        if _HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"[pdfplumber] Error: {e}")

        if not text.strip() and _HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                print(f"[PyMuPDF] Error: {e}")

        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        text = ""
        if not _HAS_DOCX:
            return text
        try:
            doc = python_docx.Document(BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"[python-docx] Error: {e}")
        return text

    # ── Structured Parsing ────────────────────────────────────────────────────

    @staticmethod
    def parse_to_structured(text: str) -> StructuredCV:
        sections = _split_into_sections(text)
        header_text = sections.get("header", sections.get("contact", ""))

        personal_info = _parse_personal_info(header_text, text)
        summary       = sections.get("summary", "").strip() or None
        skills        = _parse_skills(sections.get("skills", ""))
        experience    = _parse_experience(sections.get("experience", ""))
        education     = _parse_education(sections.get("education", ""))
        projects      = _parse_projects(sections.get("projects", ""))
        certifications = _parse_certifications(sections.get("certifications", ""))
        languages     = _parse_languages(sections.get("languages", ""))

        return StructuredCV(
            personal_information = personal_info,
            summary              = summary,
            skills               = skills,
            experience           = experience,
            education            = education,
            projects             = projects,
            certifications       = certifications,
            languages            = languages,
            raw_sections         = {k: v for k, v in sections.items()},
        )
